from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
import io
import datetime
import os
import openai
from PIL import Image
import requests
from io import BytesIO
import streamlit as st
import time

import streamlit as st
import leafmap.foliumap as leafmap
from html2image import Html2Image
from PIL import Image

def add_wms_layer(map, url, layer, layer_name, display):
    map.add_wms_layer(
        url, 
        layers=layer, 
        name=layer_name, 
        attribution=" ", 
        transparent=True,
        format="image/png",
        shown=display
        )

def main():
    with open("main.css") as f:
        st.markdown("<style>{}</style>".format(f.read()), unsafe_allow_html=True)

    st.title("Demo av funksjonalitet for automatiserte rapporter")
    #--
    c1, c2 = st.columns(2)
    with c1:
        selected_zoom = st.number_input("Zoom?", min_value = 2, value=13, max_value=20, step=1)
    with c2:
        selected_display = st.selectbox("Velg visning", ["Oversiktskart", "Løsmasserelatert", "Berggrunnsrelatert"])
    #--
    m = leafmap.Map(
        center=(59.910155, 10.811134), 
        zoom=selected_zoom,draw_control=False,
        measure_control=False,
        fullscreen_control=False,
        attribution_control=False,
        google_map="ROADMAP",
        shown=True
        )
    #--
    wms_url_list = [
        "https://geo.ngu.no/mapserver/LosmasserWMS?request=GetCapabilities&service=WMS",
        "https://geo.ngu.no/mapserver/MarinGrenseWMS4?REQUEST=GetCapabilities&SERVICE=WMS",
        "https://geo.ngu.no/mapserver/GranadaWMS5?request=GetCapabilities&service=WMS",
        "https://geo.ngu.no/geoserver/nadag/ows?request=GetCapabilities&service=WMS",
        "https://geo.ngu.no/mapserver/BerggrunnWMS3?request=GetCapabilities&SERVICE=WMS",
        "https://geo.ngu.no/mapserver/BerggrunnWMS3?request=GetCapabilities&SERVICE=WMS",
        "https://geo.ngu.no/mapserver/BerggrunnWMS3?request=GetCapabilities&SERVICE=WMS",
        
    ]
    wms_layer_list = [
        "Losmasse_flate",
        "Marin_grense_linjer",
        "Energibronn",
        "GBU_metode",
        "Berggrunn_lokal_hovedbergarter",
        "Berggrunn_regional_hovedbergarter",
        "Berggrunn_nasjonal_hovedbergarter",
    ]
    wms_name_list = [
        "Løsmasser",
        "Marin grense",            
        "Energibrønner",
        "Grunnundersøkelser",
        "Lokal berggrunn",
        "Regional berggrunn",
        "Nasjonal berggrunn",
    ]
    for i in range(0, len(wms_url_list)):
        display = False
        if selected_display == "Løsmasserelatert" and i < 4:
            display = True 
        if selected_display == "Berggrunnsrelatert" and i == 4:
            display = True
        add_wms_layer(
            m,
            wms_url_list[i],
            wms_layer_list[i],
            wms_name_list[i],
            display
        )
    m.to_streamlit(700, 400)
    m.to_html("kart.html")
    #hti = Html2Image()
    #hti.screenshot(html_file='kart.html', save_as='kart.png', size=(500, 500))
    #--
    st.header("Innhenting av data")
    #st.subheader("Eksempel: OpenAI API")
    #command = st.text_input("Lag et bilde med AI (skriv inn tekst)", value="three dogs playing chess, oil painting")
    #openai.api_key = "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"
    #number_of_images = 1
    #image_response = openai.Image.create(prompt=command, n=number_of_images, size="512x512", response_format="url")
    #for i in range(0, number_of_images):
    #    url = image_response["data"][i]["url"]
    #    response = requests.get(url)
    #    img = Image.open(BytesIO(response.content))
    #    st.image(img)
    #--
    st.subheader("Eksempel: Parametere")
    st.caption("Kan legges inn via nettsiden eller for eksempel hentes inn fra kart API / Sharepoint API")
    with st.form("Input"):
        forfatter = st.text_input("Forfatter", value="Ola Nordmann")
        oppdragsleder = st.text_input("Oppdragsleder", value="Kari Nordmann")
        oppdragsgiver = st.text_input("Oppdragsgiver", value = "Firma AS")
        oppdragsnummer = st.text_input("Oppdragsnummer", value = "635960-01")
        sted = st.text_input("Sted", value = "Trondheim")
        #--
        depth_to_bedrock = st.number_input("Dybde til fjell [m]", value=5, step=1)
        loose_material = st.selectbox("Hva slags løsmasser?", options=["hav- og fjordavsetning", "elveavsetning", "breelvavsetning", "morene"])
        st.form_submit_button("Gi input")
    #--
    if depth_to_bedrock > 15:
        setning_dybde_til_fjell = "Siden dybde til fjell > 15 m, må det gjøres supplerende grunnundersøkelser. "
    else:
        setning_dybde_til_fjell = ""
    report_text_1 = f"""Det skal gjøres geoteknisk vurdering for {sted}. Dybde til fjell var {depth_to_bedrock} m. 
    {setning_dybde_til_fjell}Løsmassene er kartlagt som {loose_material}. Viktige problemstillinger innen geoteknikk er vurdering av fundamenters bæreevne (lastkapasitet) og deres setning under belastning, jordtrykk mot støtte- og kjellermurer, stabilitet av veiskjæringer og naturlige skråninger, fundamentering av marine konstruksjoner og rørledninger. """
    #--
    document = Document("Notatmal.docx")
    styles = document.styles
    style = styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)
    document.paragraphs[0].text = f"Oppdragsgiver: {oppdragsgiver}"
    document.paragraphs[1].text = f"Oppdragsnavn: Geoteknisk rapport - {sted}"
    document.paragraphs[2].text = f"Oppdragsnummer: {oppdragsnummer} - {sted}"
    document.paragraphs[3].text = f"Utarbeidet av: {forfatter}"
    document.paragraphs[4].text = f"Oppdragsleder: {oppdragsleder}"
    document.paragraphs[5].text = f"Dato: {datetime.date.today()}"
    document.paragraphs[6].text = f"Tilgjenglighet: Åpen"

    document.paragraphs[7].text = f"Geoteknisk notat - {sted}"

    document.add_heading("Innledning", 1)
    document.add_paragraph(report_text_1)

    document.add_picture("kart.png")
    #--
    st.markdown("---")
    bio = io.BytesIO()
    document.save(bio)
    if document:
        st.download_button(
            label="Last ned rapport!",
            data=bio.getvalue(),
            file_name="Rapport.docx",
            mime="docx")

    #--
main()